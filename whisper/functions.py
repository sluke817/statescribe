import whisper
import boto3
from botocore.exceptions import ClientError
import os
import sys
import logging
import openai
from docx import Document
sys.path.append("../database")
from database import db_functions

STORAGE_PATH = "./tmp"
logging.basicConfig(filename ='whisper.log', level = logging.ERROR)


# transcribed a video from an s3 bucket and returns the text string
def transcribe_video(bucket_name, key):
    model = whisper.load_model("large", download_root="/tmp")
    s3_client = boto3.client("s3")

    logging.info("Downloading video...")
    s3_client.download_file(Bucket=bucket_name, Key=key, Filename=f"{STORAGE_PATH}/{key}")
    logging.info("Complete!")

    logging.info("Deleting video from S3 bucket...")
    s3_client.delete_object(Bucket=bucket_name, Key=key)
    logging.info("Complete!")

    logging.info("Transcribing video...")
    result = model.transcribe(f"{STORAGE_PATH}/{key}")
    logging.info("Complete!")

    os.remove(f"{STORAGE_PATH}/{key}")

    return result["text"]


def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def sentiment_analysis(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def summarize_transcription(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


def export_to_docx(summarization, filename, bucket_name=None):
    doc = Document()
    for key, value in summarization.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

    if(bucket_name != None):
        s3_client = boto3.client("s3")

        logging.info("Uploading doc...")
        s3_client.download_file(Bucket=bucket_name, Key=key, Filename=f"{STORAGE_PATH}/{key}")
        logging.info("Complete!")


# transcribed all videos from an s3 bucket and stores the values in a database
def transcribe_bucket(db, bucket_name):
    s3_client = boto3.client("s3")
    objects = s3_client.list_objects_v2(Bucket=bucket_name)
    state = bucket_name.split(".")[0]

    for obj in objects["Contents"]:
        logging.info(f"*** Starting transcription for video {obj['key']} ***")
        transcript = transcribe_video(bucket_name, obj["Key"])
        logging.info("Complete!")
        logging.info("Summarizing video...")
        summary = summarize_transcription(transcript)
        logging.info("Complete!")
        logging.info("Inserting into database...")
        sql = f"INSERT INTO hearings (state, videoID, transcript, summary, keyPoints, actionItems, sentiment) \
            VALUES ({state}, {obj['Key']}, {transcript}, {summary['abstract_summary']}, {summary['key_points']}, \
                {summary['action_items']}, {summary['sentiment']})"
        response = db_functions.query_db(db, sql, True)
        logging.info(f"DB Response: {response}")


if __name__ == "__main__":

    logging.info(f"Bucket: {sys.argv[1]}")
    
    logging.info("Connectiong to db...")
    db = db_functions.connect_to_db()
    logging.info("Complete!")

    logging.info("Transcribing bucket...")
    transcribe_bucket(db, sys.argv[1])
    logging.info("Complete!")

    logging.info("Summ...")
